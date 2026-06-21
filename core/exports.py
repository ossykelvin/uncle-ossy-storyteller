from pathlib import Path
import html
import markdown as md
from core.storage import project_dir
from core.utils import slugify, now_iso


def manuscript_markdown(project: dict, outline: dict, chapters: list, include_bible: bool = False, story_bible: dict | None = None) -> str:
    parts = [f"# {project.get('title', 'Untitled')}\n"]
    if outline.get("content"):
        parts.append("## Outline\n")
        parts.append(outline.get("content", ""))
    for c in sorted(chapters, key=lambda x: x.get("chapter_number", 0)):
        parts.append(f"\n\n## Chapter {c.get('chapter_number')}: {c.get('title', '')}\n")
        parts.append(c.get("content", ""))
    if include_bible and story_bible:
        parts.append("\n\n# Story Bible\n")
        parts.append("```json\n" + __import__('json').dumps(story_bible, indent=2, ensure_ascii=False) + "\n```")
    return "\n".join(parts)


def export_markdown(bundle: dict) -> Path:
    p = project_dir(bundle["project"]["id"]) / "exports"
    p.mkdir(exist_ok=True)
    path = p / f"{slugify(bundle['project'].get('title','project'))}.md"
    path.write_text(manuscript_markdown(bundle["project"], bundle["outline"], bundle["chapters"], True, bundle["story_bible"]), encoding="utf-8")
    return path


def export_html(bundle: dict) -> Path:
    p = project_dir(bundle["project"]["id"]) / "exports"
    p.mkdir(exist_ok=True)
    markdown_text = manuscript_markdown(bundle["project"], bundle["outline"], bundle["chapters"], True, bundle["story_bible"])
    body = md.markdown(markdown_text, extensions=["tables", "fenced_code"])
    title = html.escape(bundle["project"].get("title", "Untitled"))
    doc = f"""<!doctype html><html><head><meta charset='utf-8'><title>{title}</title>
<style>body{{font-family:Georgia,serif;max-width:900px;margin:40px auto;line-height:1.7;color:#111827}}h1,h2{{color:#0D1B3D}}pre{{background:#F2F4F7;padding:16px;overflow:auto}}</style>
</head><body>{body}</body></html>"""
    path = p / f"{slugify(bundle['project'].get('title','project'))}.html"
    path.write_text(doc, encoding="utf-8")
    return path


def export_docx(bundle: dict) -> Path:
    from docx import Document
    p = project_dir(bundle["project"]["id"]) / "exports"
    p.mkdir(exist_ok=True)
    path = p / f"{slugify(bundle['project'].get('title','project'))}.docx"
    doc = Document()
    doc.add_heading(bundle["project"].get("title", "Untitled"), 0)
    if bundle["outline"].get("content"):
        doc.add_heading("Outline", level=1)
        doc.add_paragraph(bundle["outline"].get("content", ""))
    for c in sorted(bundle["chapters"], key=lambda x: x.get("chapter_number", 0)):
        doc.add_heading(f"Chapter {c.get('chapter_number')}: {c.get('title','')}", level=1)
        for para in c.get("content", "").split("\n\n"):
            doc.add_paragraph(para)
    doc.save(path)
    return path


def export_pdf(bundle: dict) -> Path:
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.pagesizes import A4
    p = project_dir(bundle["project"]["id"]) / "exports"
    p.mkdir(exist_ok=True)
    path = p / f"{slugify(bundle['project'].get('title','project'))}.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=A4)
    styles = getSampleStyleSheet()
    story = [Paragraph(bundle["project"].get("title", "Untitled"), styles["Title"]), Spacer(1, 12)]
    for c in sorted(bundle["chapters"], key=lambda x: x.get("chapter_number", 0)):
        story.append(Paragraph(f"Chapter {c.get('chapter_number')}: {c.get('title','')}", styles["Heading1"]))
        for para in c.get("content", "").split("\n\n"):
            story.append(Paragraph(html.escape(para).replace("\n", "<br/>"), styles["BodyText"]))
            story.append(Spacer(1, 8))
    doc.build(story)
    return path


def export_epub(bundle: dict) -> Path:
    from ebooklib import epub
    p = project_dir(bundle["project"]["id"]) / "exports"
    p.mkdir(exist_ok=True)
    path = p / f"{slugify(bundle['project'].get('title','project'))}.epub"
    book = epub.EpubBook()
    book.set_identifier(bundle["project"].get("id", "uncle-ossy-project"))
    book.set_title(bundle["project"].get("title", "Untitled"))
    book.set_language("en")
    chapters = []
    for c in sorted(bundle["chapters"], key=lambda x: x.get("chapter_number", 0)):
        ch = epub.EpubHtml(title=f"Chapter {c.get('chapter_number')}", file_name=f"chap_{c.get('chapter_number')}.xhtml", lang="en")
        content = md.markdown(c.get("content", ""))
        ch.content = f"<h1>Chapter {c.get('chapter_number')}: {html.escape(c.get('title',''))}</h1>{content}"
        book.add_item(ch)
        chapters.append(ch)
    book.toc = tuple(chapters)
    book.spine = ["nav"] + chapters
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    epub.write_epub(str(path), book)
    return path
